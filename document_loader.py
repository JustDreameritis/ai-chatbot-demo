"""
document_loader.py — Multi-format document ingestion

Supports PDF, plain text, Markdown, and DOCX files.
Each loaded document is returned as a list of page dicts with text and metadata.
"""

import io
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

import chardet


@dataclass
class DocumentPage:
    """A single page or logical section from a loaded document."""

    text: str
    metadata: dict = field(default_factory=dict)

    def __post_init__(self) -> None:
        self.text = self.text.strip()


def load_document(file_path: str, original_filename: Optional[str] = None) -> List[DocumentPage]:
    """Dispatch to the correct loader based on file extension.

    Args:
        file_path: Absolute path to the file on disk.
        original_filename: Display name to embed in metadata (uses basename if None).

    Returns:
        List of DocumentPage objects, one per page / logical section.

    Raises:
        ValueError: If the file extension is not supported.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    display_name = original_filename or path.name
    ext = path.suffix.lower()

    if ext == ".pdf":
        return _load_pdf(path, display_name)
    elif ext in (".txt", ".md", ".markdown"):
        return _load_text(path, display_name)
    elif ext in (".docx",):
        return _load_docx(path, display_name)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. "
            "Supported types: PDF, TXT, MD, DOCX"
        )


def _load_pdf(path: Path, display_name: str) -> List[DocumentPage]:
    """Load a PDF file, returning one DocumentPage per page.

    Args:
        path: Path to the PDF file.
        display_name: Human-readable filename for metadata.

    Returns:
        List of DocumentPage objects.
    """
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required for PDF support: pip install PyPDF2")

    pages: List[DocumentPage] = []
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        total_pages = len(reader.pages)

        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if not text.strip():
                continue  # Skip blank pages

            pages.append(
                DocumentPage(
                    text=text,
                    metadata={
                        "source": display_name,
                        "page": i + 1,
                        "total_pages": total_pages,
                        "file_type": "pdf",
                    },
                )
            )

    if not pages:
        raise ValueError(f"No extractable text found in PDF: {display_name}")

    return pages


def _load_text(path: Path, display_name: str) -> List[DocumentPage]:
    """Load a plain text or Markdown file as a single DocumentPage.

    Auto-detects encoding using chardet.

    Args:
        path: Path to the text file.
        display_name: Human-readable filename for metadata.

    Returns:
        List with a single DocumentPage.
    """
    raw = path.read_bytes()

    # Detect encoding
    detection = chardet.detect(raw)
    encoding = detection.get("encoding") or "utf-8"

    try:
        text = raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        text = raw.decode("utf-8", errors="replace")

    file_type = "markdown" if path.suffix.lower() in (".md", ".markdown") else "text"

    return [
        DocumentPage(
            text=text,
            metadata={
                "source": display_name,
                "page": 1,
                "total_pages": 1,
                "file_type": file_type,
                "encoding": encoding,
                "size_bytes": len(raw),
            },
        )
    ]


def _load_docx(path: Path, display_name: str) -> List[DocumentPage]:
    """Load a DOCX file, returning one DocumentPage per section.

    Paragraphs are grouped into logical sections (split on heading styles).

    Args:
        path: Path to the DOCX file.
        display_name: Human-readable filename for metadata.

    Returns:
        List of DocumentPage objects (one per section or full doc if no headings).
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support: pip install python-docx"
        )

    doc = Document(str(path))

    sections: List[str] = []
    current_section: List[str] = []

    HEADING_STYLES = {
        "Heading 1", "Heading 2", "Heading 3",
        "heading 1", "heading 2", "heading 3",
    }

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name in HEADING_STYLES and current_section:
            sections.append("\n".join(current_section))
            current_section = [text]
        else:
            current_section.append(text)

    if current_section:
        sections.append("\n".join(current_section))

    # Fall back: treat entire doc as one section
    if not sections:
        full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        if not full_text:
            raise ValueError(f"No extractable text found in DOCX: {display_name}")
        sections = [full_text]

    total = len(sections)
    return [
        DocumentPage(
            text=section,
            metadata={
                "source": display_name,
                "page": i + 1,
                "total_pages": total,
                "file_type": "docx",
            },
        )
        for i, section in enumerate(sections)
    ]


def load_from_bytes(
    file_bytes: bytes, filename: str
) -> List[DocumentPage]:
    """Load a document from raw bytes (e.g. from Streamlit UploadedFile).

    Writes to a temporary file, loads it, then cleans up.

    Args:
        file_bytes: Raw file content.
        filename: Original filename (used to detect extension and for metadata).

    Returns:
        List of DocumentPage objects.
    """
    import tempfile

    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        pages = load_document(tmp_path, original_filename=filename)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    return pages
